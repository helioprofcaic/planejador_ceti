import streamlit as st
import json
import os
import sys

# Adiciona o diretório raiz do projeto (onde este utils.py está) ao sys.path
# Isso garante que módulos em pastas como 'tools' possam ser importados de qualquer lugar.
# --- NOVA ABORDAGEM DE PATH ---
# Adiciona o diretório raiz E o diretório de storage diretamente ao path
# para resolver problemas de importação em ambientes como o Streamlit Cloud.
project_root = os.path.dirname(os.path.abspath(__file__))
storage_path = os.path.join(project_root, 'tools', 'storage')

if storage_path not in sys.path:
    sys.path.insert(0, storage_path)
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from fpdf import FPDF
import io
from pypdf import PdfWriter, PdfReader
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import re

# --- ReportLab Imports ---
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# --- Integração com Google Drive ---
try:
    import google_storage
    HAS_GOOGLE_STORAGE = True
except ImportError:
    HAS_GOOGLE_STORAGE = False
    
# --- Integração com Supabase ---
try:
    import database as db
    HAS_SUPABASE = True
except (ImportError, ModuleNotFoundError) as e:
    # Tenta fallback para importação relativa ou falha silenciosamente
    try:
        from tools.storage import database as db
        HAS_SUPABASE = True
    except (ImportError, ModuleNotFoundError):
        print(f"Debug: Falha na importação do Supabase: {e}")
        HAS_SUPABASE = False

# Prioridade: Supabase > Cloud > Local
USE_SUPABASE = False
if HAS_SUPABASE:
    try:
        USE_SUPABASE = (
            st.secrets.get("supabase", {}).get("usar_supabase", False) and
            db.is_db_connected()
        )
    except Exception:
        USE_SUPABASE = False

# Se Supabase estiver ativo, desativa o Cloud Storage para evitar conflitos
USE_CLOUD_STORAGE = False
if not USE_SUPABASE:
    try:
        if os.environ.get("FORCE_LOCAL_MODE") == "1":
            USE_CLOUD_STORAGE = False
        else:
            USE_CLOUD_STORAGE = (
                HAS_GOOGLE_STORAGE and
                st.secrets.get("drive", {}).get("usar_nuvem", False)
            )
    except Exception:
        USE_CLOUD_STORAGE = False

def aplicar_estilo():
    """Aplica o CSS global baseado nas configurações de sessão."""
    tema = st.session_state.get('tema', "Padrão")
    tamanho_fonte = st.session_state.get('tamanho_fonte', 16)
    
    padding_top = "0rem" if tema == "Compacto" else "2rem"
    font_style = "Arial Narrow" if tema == "Compacto" else "sans-serif"
    
    st.markdown(f"""
        <style>
        html, body {{
            font-size: {tamanho_fonte}px;
            font-family: {font_style};
        }}
        [class*="st-"] {{
            font-size: {tamanho_fonte}px;
        }}
        /* Correção global para setas do st.expander aparecendo como texto */
        [data-testid="stExpander"] summary::after, [data-testid="stExpander"] summary::before {{
            content: "" !important;
        }}
        .main .block-container {{
            padding-top: {padding_top};
        }}
        /* Ajuste para as tabelas não ficarem gigantes */
        .stDataFrame div[data-testid="stTable"] {{
            font-size: {tamanho_fonte - 2}px;
        }}
        
        /* Ocultar elementos de UI do Streamlit que são em inglês */
        /* #MainMenu {{visibility: hidden;}} */
        footer {{visibility: hidden;}}
        /* [data-testid="stToolbar"] {{visibility: hidden;}} */

        /* Correção para Mobile: Forçar exibição do botão de menu */
        @media (max-width: 768px) {{
            [data-testid="stSidebarCollapsedControl"] {{
                visibility: visible !important;
                display: block !important;
            }}
        }}
        </style>
        """, unsafe_allow_html=True)

def carregar_dados():
    """[DEPRECATED] Carrega o arquivo ementas.json da pasta data."""
    caminho = os.path.join("data", "ementas.json")
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return {}

def carregar_ementas_oficiais():
    """[DEPRECATED] Carrega o arquivo ementas_oficiais.json da pasta data."""
    caminho = os.path.join("data", "ementas_oficiais.json")
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return {}

def carregar_ementas_trimestre():
    """[DEPRECATED] Carrega o arquivo ementas_geral_1trimestre.json da pasta data."""
    caminho = os.path.join("data", "ementas_geral_1trimestre.json")
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return {}

def carregar_escola_db():
    """Carrega o banco de dados da escola (escola_db.json) do local ou da nuvem."""
    filename = "escola_db.json"
    default_data = {"turmas": {}, "professores": []}
    
    data = default_data
    if USE_CLOUD_STORAGE:
        data = google_storage.load_json(filename, default_value=default_data, folder_path=['data', 'escola'])
    else:
        caminho = os.path.join("data", "escola", filename)
        if os.path.exists(caminho):
            try:
                with open(caminho, "r", encoding="utf-8-sig") as f:
                    content = f.read()
                    if content:
                        data = json.loads(content)
            except json.JSONDecodeError:
                st.warning(f"⚠️ Arquivo `{filename}` está mal formatado. Usando dados padrão.")
                data = default_data
    
    # --- AUTO-CORREÇÃO: Sincronizar com alunos.json se não houver turmas ---
    # Se o escola_db não tiver turmas, tenta pegar do alunos.json
    if not data.get("turmas"):
        alunos = carregar_alunos()
        if alunos:
            turmas_novas = {}
            for turma in alunos.keys():
                turmas_novas[turma] = {"componentes": []}
            data["turmas"] = turmas_novas
            # Opcional: Salvar essa inferência de volta para persistir
            # salvar_escola_db(data) 
            
    return data

def salvar_escola_db(dados):
    """Salva o arquivo escola_db.json no local ou na nuvem."""
    filename = "escola_db.json"
    
    if USE_CLOUD_STORAGE:
        google_storage.save_json(filename, dados, folder_path=['data', 'escola'])
    else:
        caminho = os.path.join("data", "escola", filename)
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(dados, f, indent=2, ensure_ascii=False)

def carregar_calendario_letivo():
    """Carrega o arquivo calendario_letivo_2026.json da pasta data."""
    caminho = os.path.join("data", "calendario_letivo_2026.json")
    # Padrão de fallback (Bússola do Tempo padrão)
    padrao = {
        "trimestres": {
            "1º": {"inicio": "2026-02-19", "fim": "2026-05-22", "semana_inicio": 0, "semana_fim": 13},
            "2º": {"inicio": "2026-05-25", "fim": "2026-08-28", "semana_inicio": 13, "semana_fim": 26},
            "3º": {"inicio": "2026-08-31", "fim": "2026-12-18", "semana_inicio": 26, "semana_fim": 40}
        }
    }
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            try:
                return json.load(f)
            except:
                return padrao
    return padrao

def carregar_curriculo_db():
    """Carrega o banco de dados do currículo (curriculo_db.json) do local ou da nuvem."""
    filename = "curriculo_db.json"
    default_data = {"BASICO": {}, "APROFUNDAMENTO": {}, "EPT": {}}
    
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value=default_data)
        
    caminho = os.path.join("data", filename)
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return default_data

def carregar_habilidades_csv():
    """Carrega habilidades de arquivos CSV na pasta data."""
    dados_csv = {}
    caminho_dir = "data"
    if not os.path.exists(caminho_dir):
        return dados_csv
        
    arquivos = [f for f in os.listdir(caminho_dir) if f.startswith("habilidades") and f.endswith(".csv")]
    
    for arquivo in arquivos:
        try:
            df = pd.read_csv(os.path.join(caminho_dir, arquivo))
            # Normaliza colunas para minúsculas e sem espaços
            df.columns = [c.lower().strip() for c in df.columns]
            
            if 'componente' in df.columns:
                for _, row in df.iterrows():
                    comp = str(row['componente']).strip()
                    
                    if comp not in dados_csv:
                        dados_csv[comp] = {"competencia": "", "habilidades": [], "objetos": []}
                    
                    # Preenche competência (pega a primeira não nula encontrada)
                    if 'competencia' in df.columns and pd.notna(row['competencia']) and not dados_csv[comp]["competencia"]:
                        dados_csv[comp]["competencia"] = row['competencia']
                        
                    if 'habilidade' in df.columns and pd.notna(row['habilidade']):
                        h = row['habilidade']
                        if h not in dados_csv[comp]["habilidades"]:
                            dados_csv[comp]["habilidades"].append(h)
                            
                    col_obj = 'objeto_conhecimento' if 'objeto_conhecimento' in df.columns else 'conteudo'
                    if col_obj in df.columns and pd.notna(row[col_obj]):
                        o = row[col_obj]
                        if o not in dados_csv[comp]["objetos"]:
                            dados_csv[comp]["objetos"].append(o)
        except Exception as e:
            print(f"Erro ao ler {arquivo}: {e}")
            
    return dados_csv

def salvar_ementas_trimestre(dados):
    """[DEPRECATED] Salva o arquivo ementas_geral_1trimestre.json na pasta data."""
    caminho = os.path.join("data", "ementas_geral_1trimestre.json")
    os.makedirs("data", exist_ok=True)
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(dados, f, indent=2, ensure_ascii=False)

def carregar_config_componentes():
    """Carrega o arquivo config_componentes.json da pasta data (local ou nuvem)."""
    filename = "config_componentes.json"
    default_data = {
        "MAPEAMENTO_POR_CHAVE": {},
        "PADRAO_GERAL": {"tipo_curso": "Anual / Regular", "duracao_semanas": 13},
        "PADRAO_TECNICO_MODULAR": {"tipo_curso": "Modular Mensal (40h)", "duracao_semanas": 5}
    }
    
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value=default_data)
        
    caminho = os.path.join("data", filename)
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return default_data

def salvar_config_componentes(dados):
    """Salva o arquivo config_componentes.json na pasta data (local ou nuvem)."""
    filename = "config_componentes.json"
    if USE_CLOUD_STORAGE:
        google_storage.save_json(filename, dados)
    else:
        caminho = os.path.join("data", filename)
        os.makedirs("data", exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(dados, f, indent=2, ensure_ascii=False)

def salvar_planejamento(dados):
    """Salva um planejamento específico em data/planejamentos.json (local ou nuvem)."""
    filename = "planejamentos.json"
    
    # Carrega os planejamentos existentes ou cria um novo dicionário
    if USE_CLOUD_STORAGE:
        todos = google_storage.load_json(filename, default_value={})
    else:
        caminho = os.path.join("data", filename)
        if os.path.exists(caminho):
            try:
                with open(caminho, "r", encoding="utf-8-sig") as f:
                    todos = json.load(f)
            except json.JSONDecodeError:
                todos = {}
        else:
            todos = {}
        
    # Chave única para identificar o plano
    trimestre = dados.get("trimestre", "1º")
    chave = f"{dados['turma']}_{dados['componente']}_{dados['escala']}_{trimestre}"
    todos[chave] = dados

    # Salva o arquivo atualizado
    if USE_CLOUD_STORAGE:
        google_storage.save_json(filename, todos)
    else:
        caminho = os.path.join("data", filename)
        os.makedirs("data", exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(todos, f, indent=2, ensure_ascii=False)

def carregar_planejamento(turma, componente, escala, trimestre="1º"):
    """Carrega um planejamento específico se existir (local ou nuvem)."""
    filename = "planejamentos.json"
    
    if USE_CLOUD_STORAGE:
        todos = google_storage.load_json(filename, default_value={})
    else:
        caminho = os.path.join("data", filename)
        if os.path.exists(caminho):
            try:
                with open(caminho, "r", encoding="utf-8-sig") as f:
                    todos = json.load(f)
            except json.JSONDecodeError:
                todos = {}
        else:
            return None

    chave = f"{turma}_{componente}_{escala}_{trimestre}"
    return todos.get(chave)

def carregar_alunos():
    """Carrega o arquivo alunos.json da pasta data (local ou nuvem)."""
    # --- PROTEÇÃO DE DADOS ---
    # Visitantes não veem a lista de alunos
    if st.session_state.get("professor") == "Visitante":
        return {}

    filename = "alunos.json"
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value={}, folder_path=['data', 'escola'])
        
    caminho = os.path.join("data", "escola", filename)
    if os.path.exists(caminho):
        try:
            with open(caminho, "r", encoding="utf-8-sig") as f:
                content = f.read()
                if content:
                    return json.loads(content)
        except json.JSONDecodeError:
            st.warning(f"⚠️ Arquivo `{filename}` está mal formatado. Nenhum aluno carregado.")
            return {}
    return {}

def salvar_alunos(dados):
    """Salva o arquivo alunos.json na pasta data (local ou nuvem)."""
    filename = "alunos.json"
    if USE_CLOUD_STORAGE:
        return google_storage.save_json(filename, dados, folder_path=['data', 'escola'])
    else:
        caminho = os.path.join("data", "escola", filename)
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(dados, f, indent=2, ensure_ascii=False)
        return True

def salvar_dados_json(caminho_arquivo, dados_df):
    """Salva um DataFrame em um arquivo JSON (local ou nuvem)."""
    filename = os.path.basename(caminho_arquivo)
    
    if USE_CLOUD_STORAGE:
        # Converte DataFrame para lista de dicionários para ser compatível com JSON
        dados_dict = dados_df.to_dict(orient='records')
        
        # Verifica se deve salvar na subpasta 'perfis'
        folder_path = ['data']
        path_parts = os.path.normpath(caminho_arquivo).split(os.sep)
        if 'perfis' in path_parts:
            folder_path = ['data', 'perfis']
        elif 'frequencia' in path_parts:
            folder_path = ['data', 'frequencia']
        elif 'avaliacoes' in path_parts:
            folder_path = ['data', 'avaliacoes']
            
        google_storage.save_json(filename, dados_dict, folder_path=folder_path)
    else:
        # Salva localmente
        os.makedirs(os.path.dirname(caminho_arquivo), exist_ok=True)
        dados_df.to_json(caminho_arquivo, orient='records', indent=4, force_ascii=False)

def carregar_dados_json(caminho_arquivo):
    """Carrega um DataFrame de um arquivo JSON (local ou nuvem)."""
    filename = os.path.basename(caminho_arquivo)

    if USE_CLOUD_STORAGE:
        # Usa um sentinela para distinguir "arquivo não encontrado" de "arquivo vazio"
        sentinela = {"__arquivo_nao_encontrado__": True}
        
        # Verifica se deve carregar da subpasta 'perfis'
        folder_path = ['data']
        path_parts = os.path.normpath(caminho_arquivo).split(os.sep)
        if 'perfis' in path_parts:
            folder_path = ['data', 'perfis']
        elif 'frequencia' in path_parts:
            folder_path = ['data', 'frequencia']
        elif 'avaliacoes' in path_parts:
            folder_path = ['data', 'avaliacoes']
            
        dados_dict = google_storage.load_json(filename, default_value=sentinela, folder_path=folder_path)
        
        if dados_dict == sentinela:
            return None
            
        if dados_dict is not None:
            # Se o arquivo existe na nuvem mas está vazio, retorna um DF vazio
            if not dados_dict:
                return pd.DataFrame()
            return pd.DataFrame(dados_dict)
        return None # Retorna None se o arquivo não existe na nuvem

    # Lógica local
    if os.path.exists(caminho_arquivo):
        if os.path.getsize(caminho_arquivo) > 0:
            try:
                with open(caminho_arquivo, "r", encoding="utf-8-sig") as f:
                    return pd.DataFrame(json.load(f))
            except (ValueError, json.JSONDecodeError):
                print(f"Aviso: Arquivo JSON local inválido em {caminho_arquivo}.")
                return None
    return None

def listar_arquivos_dados(prefixo, subfolder=None):
    """Lista arquivos de dados (frequencia, qualitativo) locais ou na nuvem."""
    arquivos = []
    if USE_CLOUD_STORAGE:
        folder_path = ['data']
        if subfolder:
            folder_path.append(subfolder)
            
        todos_arquivos = google_storage.list_files_in_path(folder_path)
        arquivos = [f['name'] for f in todos_arquivos if prefixo in f['name']]
    else:
        search_path = "data"
        if subfolder:
            search_path = os.path.join("data", subfolder)
            
        if os.path.exists(search_path):
            arquivos = [f for f in os.listdir(search_path) if f.startswith(prefixo) and f.endswith(".json")]
            
    return arquivos

def listar_subpastas(caminho_relativo_list):
    """Lista subpastas de um caminho (local ou nuvem)."""
    if USE_CLOUD_STORAGE:
        # list_files_in_path retorna uma lista de dicts [{'id': ..., 'name': ...}]
        folders = google_storage.list_files_in_path(caminho_relativo_list, mime_type='application/vnd.google-apps.folder')
        return sorted([f['name'] for f in folders])
    else:
        # Modo Local
        caminho_completo = os.path.join(*caminho_relativo_list)
        if not os.path.exists(caminho_completo):
            return []
        return sorted([d for d in os.listdir(caminho_completo) if os.path.isdir(os.path.join(caminho_completo, d))])

def listar_arquivos_md(caminho_relativo_list):
    """Lista arquivos .md de um caminho (local ou nuvem)."""
    if USE_CLOUD_STORAGE:
        files = google_storage.list_files_in_path(caminho_relativo_list)
        # Filtra por nome, pois o mime type pode variar (text/markdown, application/octet-stream)
        return sorted([f['name'] for f in files if f['name'].lower().endswith('.md')])
    else:
        caminho_completo = os.path.join(*caminho_relativo_list)
        if not os.path.exists(caminho_completo):
            return []
        return sorted([f for f in os.listdir(caminho_completo) if f.endswith(".md")])

def carregar_arquivo_texto(caminho_relativo):
    """Carrega um arquivo de texto (como .md) do local ou da nuvem."""
    if USE_CLOUD_STORAGE:
        filename = os.path.basename(caminho_relativo)
        # 'data/aulas/turma/disciplina/arquivo.md' -> ['data', 'aulas', 'turma', 'disciplina']
        folder_path_parts = os.path.normpath(os.path.dirname(caminho_relativo)).split(os.sep)
        return google_storage.load_text(filename, folder_path=folder_path_parts, default_value="")
    else:
        if os.path.exists(caminho_relativo):
            with open(caminho_relativo, "r", encoding="utf-8") as f:
                return f.read()
    return ""

def salvar_arquivo_texto(caminho_relativo, conteudo):
    """Salva um arquivo de texto (como .md) no local ou na nuvem."""
    if USE_CLOUD_STORAGE:
        filename = os.path.basename(caminho_relativo)
        folder_path_parts = os.path.normpath(os.path.dirname(caminho_relativo)).split(os.sep)
        return google_storage.save_text(filename, conteudo, folder_path=folder_path_parts)
    else:
        # Salva localmente
        os.makedirs(os.path.dirname(caminho_relativo), exist_ok=True)
        with open(caminho_relativo, "w", encoding="utf-8") as f:
            f.write(conteudo)
        return True

def listar_pdfs_referencia():
    """Lista PDFs disponíveis na pasta 'pdf' (Nuvem) ou 'data/pdf' (Local)."""
    if USE_CLOUD_STORAGE:
        arquivos = google_storage.list_files_in_subfolder('pdf', 'application/pdf')
        return arquivos # Retorna lista de dicts [{'id':..., 'name':...}]
    else:
        # Modo Local
        caminho_pdf = os.path.join("data", "pdf")
        if not os.path.exists(caminho_pdf):
            os.makedirs(caminho_pdf, exist_ok=True)
        
        arquivos = []
        for f in os.listdir(caminho_pdf):
            if f.lower().endswith(".pdf"):
                arquivos.append({'name': f, 'id': os.path.join(caminho_pdf, f)})
        return arquivos

def extrair_texto_pdf_referencia(file_id_or_path):
    """Extrai texto de um PDF (seja do Drive ou Local)."""
    texto_completo = ""
    try:
        if USE_CLOUD_STORAGE:
            # Baixa bytes do Drive
            pdf_bytes = google_storage.download_file_bytes(file_id_or_path)
            if pdf_bytes:
                reader = PdfReader(pdf_bytes)
                for page in reader.pages:
                    texto_completo += page.extract_text() + "\n"
        else:
            # Lê local
            with open(file_id_or_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    texto_completo += page.extract_text() + "\n"
    except Exception as e:
        print(f"Erro ao ler PDF: {e}")
        return ""
        
    return texto_completo

def gerar_docx_planejamento(escola, professor, turma, componente, escala, comp_geral, df, trimestre="1º", municipio="", lista_aulas=""):
    """Gera o DOCX do planejamento escolar."""
    doc = Document()
    
    # Título
    heading = doc.add_heading('Planejamento Escolar', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações
    p = doc.add_paragraph()
    p.add_run(f"Escola: {escola}\n").bold = True
    p.add_run(f"Professor: {professor} | Turma: {turma}\n")
    p.add_run(f"Componente: {componente} | Escala: {escala} | Trimestre: {trimestre} | Município: {municipio}\n")
    
    # Competência
    doc.add_heading('Competência Geral:', level=2)
    doc.add_paragraph(comp_geral)
    
    # Lista de Aulas (Opcional)
    if lista_aulas:
        doc.add_heading('Lista de Aulas / Conteúdos:', level=2)
        doc.add_paragraph(lista_aulas)
    
    # Tabela
    if not df.empty:
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        
        # Cabeçalho
        hdr_cells = t.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        # Dados
        for _, row in df.iterrows():
            row_cells = t.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_docx_frequencia(turma, data_aula, df):
    """Gera o DOCX da lista de frequência."""
    doc = Document()
    
    heading = doc.add_heading('Lista de Frequência', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Turma: {turma}\n").bold = True
    p.add_run(f"Data: {data_aula.strftime('%d/%m/%Y')}")
    
    if not df.empty:
        # Tabela com 3 colunas: Nº, Nome, Assinatura
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        
        hdr_cells = t.rows[0].cells
        hdr_cells[0].text = "Nº"
        hdr_cells[1].text = "Nome do Aluno"
        hdr_cells[2].text = "Assinatura / Presença"
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            
        for _, row in df.iterrows():
            row_cells = t.add_row().cells
            row_cells[0].text = str(row['Nº'])
            row_cells[1].text = str(row['Nome do Aluno'])
            row_cells[2].text = "" # Espaço para assinatura
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _registrar_fontes_reportlab():
    """Garante que as fontes DejaVu existam e as registra no ReportLab."""
    font_dir = os.path.join("data", "fonts")
    font_regular = os.path.join(font_dir, "DejaVuSans.ttf")
    font_bold = os.path.join(font_dir, "DejaVuSans-Bold.ttf")
    font_italic = os.path.join(font_dir, "DejaVuSans-Oblique.ttf")
    font_bold_italic = os.path.join(font_dir, "DejaVuSans-BoldOblique.ttf")

    # --- Download automático das fontes se não existirem ---
    try:
        if not os.path.exists(font_regular) or not os.path.exists(font_bold) or not os.path.exists(font_italic) or not os.path.exists(font_bold_italic):
            import urllib.request
            os.makedirs(font_dir, exist_ok=True)
            base_url = "https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/"
            
            font_files_to_download = {
                "DejaVuSans.ttf": font_regular,
                "DejaVuSans-Bold.ttf": font_bold,
                "DejaVuSans-Oblique.ttf": font_italic,
                "DejaVuSans-BoldOblique.ttf": font_bold_italic
            }

            for filename, filepath in font_files_to_download.items():
                if not os.path.exists(filepath):
                    print(f"Baixando {filename}...")
                    urllib.request.urlretrieve(base_url + filename, filepath)
    except Exception as e:
        print(f"Aviso: Não foi possível baixar as fontes automaticamente: {e}")

    # --- Registro das fontes no ReportLab ---
    font_family = "Helvetica"
    if os.path.exists(font_regular) and os.path.exists(font_bold) and os.path.exists(font_italic) and os.path.exists(font_bold_italic):
        try:
            pdfmetrics.registerFont(TTFont('DejaVu', font_regular))
            pdfmetrics.registerFont(TTFont('DejaVu-Bold', font_bold))
            pdfmetrics.registerFont(TTFont('DejaVu-Italic', font_italic))
            pdfmetrics.registerFont(TTFont('DejaVu-BoldItalic', font_bold_italic))
            pdfmetrics.registerFontFamily('DejaVu', normal='DejaVu', bold='DejaVu-Bold', italic='DejaVu-Italic', boldItalic='DejaVu-BoldItalic')
            font_family = "DejaVu"
        except Exception as e:
            print(f"Erro ao registrar fonte DejaVu para ReportLab: {e}")
    
    return font_family

def gerar_pdf_planejamento(escola, professor, turma, componente, escala, comp_geral, df, trimestre, municipio, lista_aulas=""):
    """Gera o PDF do planejamento escolar usando ReportLab."""
    buffer = io.BytesIO()
    font_family = _registrar_fontes_reportlab()

    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=15*mm, leftMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=10, leading=12))
    styles.add(ParagraphStyle(name='BodyBold', fontName=f"{font_family}-Bold", fontSize=10, leading=12))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=16, alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle(name='H2', fontName=f"{font_family}-Bold", fontSize=12, spaceBefore=10, spaceAfter=4))
    styles.add(ParagraphStyle(name='TableHeader', fontName=f"{font_family}-Bold", fontSize=8, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_family, fontSize=8))

    story = []
    story.append(Paragraph("Planejamento Escolar", styles['H1']))
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph(f"<b>Escola:</b> {escola}", styles['Body']))
    story.append(Paragraph(f"<b>Professor:</b> {professor} | <b>Turma:</b> {turma}", styles['Body']))
    story.append(Paragraph(f"<b>Componente:</b> {componente} | <b>Escala:</b> {escala} | <b>Trimestre:</b> {trimestre} | <b>Município:</b> {municipio}", styles['Body']))
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph("Competência Geral:", styles['H2']))
    story.append(Paragraph(comp_geral.replace('\n', '<br/>'), styles['Body']))

    if lista_aulas:
        story.append(Paragraph("Lista de Aulas / Conteúdos:", styles['H2']))
        story.append(Paragraph(lista_aulas.replace('\n', '<br/>'), styles['Body']))

    if not df.empty:
        story.append(Spacer(1, 6*mm))
        
        # --- Tabela ---
        header = [Paragraph(str(col), styles['TableHeader']) for col in df.columns]
        data = [header]

        for _, row in df.iterrows():
            data_row = [Paragraph(str(row[col]), styles['TableCell']) for col in df.columns]
            data.append(data_row)

        # Pesos para distribuição de largura
        pesos = {
            "Nº": 0.5, "Mês": 0.8, "Semana": 1.2, "Aula": 0.6, "Período": 1.2,
            "Habilidade": 3.0, "Habilidades Integradas": 2.0, "Objetivos de Aprendizagem": 2.5,
            "Objeto do Conhecimento": 2.5, "Metodologia": 1.5, "Material de Apoio": 1.5,
            "Estratégia de Avaliação": 1.5
        }
        total_peso = sum(pesos.get(col, 1.5) for col in df.columns)
        available_width = doc.width
        col_widths = [(pesos.get(col, 1.5) / total_peso) * available_width for col in df.columns]

        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), f"{font_family}-Bold"),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        story.append(table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_capa_resumo(lista_planos):
    """Gera uma página de capa com o resumo dos planos na cesta usando ReportLab."""
    buffer = io.BytesIO()
    if not lista_planos:
        return buffer.getvalue()

    font_family = _registrar_fontes_reportlab()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=12, leading=14))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=16, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name='H2', fontName=f"{font_family}-Bold", fontSize=12, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name='TableHeader', fontName=f"{font_family}-Bold", fontSize=10, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_family, fontSize=9))

    story = []
    primeiro = lista_planos[0]
    escola = primeiro.get('escola', '')
    professor = primeiro.get('professor', '')

    story.append(Paragraph("Resumo do Planejamento Integrado", styles['H1']))
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph(f"<b>Escola:</b> {escola}", styles['Body']))
    story.append(Paragraph(f"<b>Professor:</b> {professor}", styles['Body']))
    story.append(Paragraph("<b>Início do Ano Letivo:</b> 19/02/2026", styles['Body']))
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph("Componentes Curriculares Listados:", styles['H2']))

    # --- Tabela de Resumo ---
    header = [
        Paragraph("Componente", styles['TableHeader']),
        Paragraph("Turma", styles['TableHeader']),
        Paragraph("Trimestre", styles['TableHeader']),
        Paragraph("Previsão Início", styles['TableHeader'])
    ]
    data = [header]

    calendario = carregar_calendario_letivo()
    trimestres_data = calendario.get("trimestres", {})

    for plano in lista_planos:
        comp = str(plano.get('componente', ''))
        turma = str(plano.get('turma', ''))
        trim = str(plano.get('trimestre', '1º'))
        
        data_inicio_str = trimestres_data.get(trim, {}).get("inicio", "2026-02-19")
        try:
            data_str = date.fromisoformat(data_inicio_str).strftime('%d/%m/%Y')
        except ValueError:
            data_str = "19/02/2026"

        data.append([
            Paragraph(comp, styles['TableCell']),
            Paragraph(turma, styles['TableCell']),
            Paragraph(trim, styles['TableCell']),
            Paragraph(data_str, styles['TableCell']),
        ])

    table = Table(data, colWidths=[100*mm, 70*mm, 30*mm, 40*mm], repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    story.append(table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def consolidar_planos(lista_planos):
    """
    Recebe uma lista de dicionários com dados dos planos e gera um único PDF.
    Cada item da lista deve conter: escola, professor, turma, componente, escala, comp_geral, df (DataFrame).
    """
    merger = PdfWriter()
    
    # Adiciona capa com resumo
    if lista_planos:
        capa_bytes = gerar_capa_resumo(lista_planos)
        if capa_bytes: merger.append(io.BytesIO(capa_bytes))
    
    for plano in lista_planos: #itera sobre os planos da cesta
        escola = plano['escola']
        professor = plano['professor']
        turma = plano['turma']
        componente = plano['componente']
        escala = plano['escala']
        comp_geral = plano['comp_geral']
        df = plano['df']
        municipio = plano.get('municipio', "")  # Retorna string vazia se não existir
        trimestre = plano.get('trimestre', '')
        lista_aulas = plano.get('lista_aulas', "")
        
        # Cria um dicionário com os argumentos esperados
        args = {
            'escola': escola, 'professor': professor, 'turma': turma,
            'componente': componente, 'escala': escala, 'comp_geral': comp_geral,
            'df': df, 'trimestre': trimestre, 'municipio': municipio, 'lista_aulas': lista_aulas
        }

        pdf_bytes = gerar_pdf_planejamento(**args)
        merger.append(io.BytesIO(pdf_bytes))
    
    output_buffer = io.BytesIO()
    merger.write(output_buffer)
    merger.close()
    output_buffer.seek(0)
    return output_buffer

def gerar_pdf_frequencia(escola, professor, turma, data_aula, df):
    """Gera o PDF da lista de frequência usando ReportLab."""
    buffer = io.BytesIO()
    font_family = _registrar_fontes_reportlab()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=15*mm, leftMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=10, leading=12))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=14, alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle(name='TableHeader', fontName=f"{font_family}-Bold", fontSize=9, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_family, fontSize=8))
    styles.add(ParagraphStyle(name='TableCellCenter', parent=styles['TableCell'], alignment=TA_CENTER))

    story = []
    story.append(Paragraph(f"Lista de Frequência - {escola}", styles['H1']))
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph(f"<b>Professor:</b> {professor} | <b>Turma:</b> {turma}", styles['Body']))
    story.append(Paragraph(f"<b>Data:</b> {data_aula.strftime('%d/%m/%Y')}", styles['Body']))
    story.append(Spacer(1, 6*mm))

    # --- Tabela de Frequência ---
    header = [
        Paragraph("Nº", styles['TableHeader']),
        Paragraph("Nome do Aluno", styles['TableHeader']),
        Paragraph("Status", styles['TableHeader']),
        Paragraph("Assinatura", styles['TableHeader']),
    ]
    data = [header]

    for _, row in df.iterrows():
        status = "P" if row.get('Presença', False) else "F"
        data.append([
            Paragraph(str(row['Nº']), styles['TableCellCenter']),
            Paragraph(str(row['Nome do Aluno']), styles['TableCell']),
            Paragraph(status, styles['TableCellCenter']),
            "" # Espaço para assinatura
        ])

    table = Table(data, colWidths=[15*mm, 95*mm, 20*mm, 50*mm], repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    story.append(table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_pdf_qualitativo(escola, professor, turma, df, componente="", contexto=""):
    """Gera o PDF da ficha qualitativa usando ReportLab."""
    buffer = io.BytesIO()
    font_family = _registrar_fontes_reportlab()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=9, leading=11))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=12, alignment=TA_CENTER, spaceAfter=4))
    styles.add(ParagraphStyle(name='TableHeader', fontName=f"{font_family}-Bold", fontSize=7, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_family, fontSize=7, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='TableCellLeft', parent=styles['TableCell'], alignment=TA_LEFT))

    story = []
    story.append(Paragraph("Ficha de Acompanhamento Qualitativo", styles['H1']))
    story.append(Spacer(1, 2*mm))
    story.append(Paragraph(f"<b>Escola:</b> {escola}", styles['Body']))
    story.append(Paragraph(f"<b>Professor:</b> {professor} | <b>Turma:</b> {turma}", styles['Body']))
    if componente or contexto:
        story.append(Paragraph(f"<b>Componente:</b> {componente} | <b>Contexto:</b> {contexto}", styles['Body']))
    story.append(Spacer(1, 4*mm))

    # --- Tabela ---
    header = [Paragraph(h, styles['TableHeader']) for h in ["Nº", "Nome do Estudante", "Particip.", "Entrega", "Autonomia", "NM1", "NM2", "NM3", "MT", "Rec.", "Final"]]
    data = [header]

    for _, row in df.iterrows():
        data.append([
            Paragraph(str(row['Nº']), styles['TableCell']),
            Paragraph(str(row['Nome do Estudante']), styles['TableCellLeft']),
            Paragraph(str(row.get('Participação', '')), styles['TableCell']),
            Paragraph(str(row.get('Entrega', '')), styles['TableCell']),
            Paragraph(str(row.get('Autonomia', '')), styles['TableCell']),
            Paragraph(str(row.get('NM1', '')) if pd.notna(row.get('NM1')) else "", styles['TableCell']),
            Paragraph(str(row.get('NM2', '')) if pd.notna(row.get('NM2')) else "", styles['TableCell']),
            Paragraph(str(row.get('NM3', '')) if pd.notna(row.get('NM3')) else "", styles['TableCell']),
            Paragraph(str(row.get('MT', '')) if pd.notna(row.get('MT')) else "", styles['TableCell']),
            Paragraph(str(row.get('Recuperação', '')) if pd.notna(row.get('Recuperação')) else "", styles['TableCell']),
            Paragraph(str(row.get('Nota Final', '')) if pd.notna(row.get('Nota Final')) else "", styles['TableCell']),
        ])

    col_widths = [10*mm, 90*mm, 18*mm, 18*mm, 18*mm, 12*mm, 12*mm, 12*mm, 12*mm, 12*mm, 15*mm]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    story.append(table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def init_db():
    """[DEPRECATED] Função mantida apenas para compatibilidade, não faz nada."""
    pass

def sincronizar_bd():
    """[DEPRECATED] Função desativada na versão Cloud."""
    return 0

def importar_alunos_db():
    """[DEPRECATED]"""
    return 0

def listar_turmas_db():
    """Lista turmas disponíveis diretamente do JSON de alunos."""
    alunos = carregar_alunos()
    if alunos and isinstance(alunos, dict):
        return sorted(list(alunos.keys()))
    return []

def listar_alunos_turma_db(turma):
    """Retorna lista de alunos de uma turma diretamente do JSON."""
    alunos = carregar_alunos()
    return alunos.get(turma, [])

def carregar_perfil_professor():
    """Carrega o perfil do professor de data/professor_config.json (local ou nuvem)."""
    filename = "professor_config.json"
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value={})
        
    caminho = os.path.join("data", filename)
    if os.path.exists(caminho):
        try:
            with open(caminho, "r", encoding="utf-8-sig") as f:
                content = f.read()
                if content:
                    return json.loads(content)
        except json.JSONDecodeError:
            return {}
    return {}

def salvar_perfil_professor(perfil):
    """Salva o perfil do professor em data/professor_config.json (local ou nuvem)."""
    filename = "professor_config.json"
    
    if USE_CLOUD_STORAGE:
        return google_storage.save_json(filename, perfil)
    else:
        caminho = os.path.join("data", filename)
        os.makedirs("data", exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(perfil, f, indent=2, ensure_ascii=False)
        return True

def listar_professores_db():
    """Lista os nomes dos professores cadastrados no escola_db.json."""
    escola_db = carregar_escola_db()
    return escola_db.get("professores", [])

def atualizar_lista_professores_db(novo_professor):
    """Adiciona um novo professor à lista geral em escola_db.json se não existir."""
    escola_db = carregar_escola_db()
    professores = escola_db.get("professores", [])
    
    # Verifica se já existe (case insensitive)
    if novo_professor.upper() not in [p.upper() for p in professores]:
        professores.append(novo_professor)
        professores.sort()
        escola_db["professores"] = professores
        salvar_escola_db(escola_db)

def verificar_senha(senha_input, tipo="admin"):
    """
    Verifica a senha informada contra as senhas salvas no perfil do administrador (Helio Lima).
    Tipos: 'admin', 'usuario', 'professor'.
    """
    # Carrega o perfil do admin (Helio Lima)
    # Se o arquivo não existir, retorna dict vazio
    perfil_admin = carregar_perfil_professor_db("Helio Lima")
    
    # Obtém o dicionário de senhas do perfil
    senhas_definidas = perfil_admin.get("senhas", {})
    
    # Garante que senhas_definidas seja um dicionário (caso venha None ou vazio de forma errada)
    if not isinstance(senhas_definidas, dict):
        senhas_definidas = {}
        
    # Recupera a senha correta. Se não existir, retorna None (acesso negado)
    senha_correta = senhas_definidas.get(tipo)
    
    # Se a senha não estiver definida no JSON, bloqueia o acesso
    if not senha_correta:
        return False
    
    return senha_input == senha_correta

def garantir_perfil_visitante():
    """Garante que o perfil de visitante exista para login padrão."""
    filename = "perfil_visitante.json"
    perfil_visitante = {
        "professor": "Visitante",
        "email": "",
        "municipio": "",
        "vinculos": []
    }
    
    if USE_CLOUD_STORAGE:
        # Verifica se existe, se não, cria
        if not google_storage.load_json(filename, silent=True, folder_path=['data', 'perfis']):
            google_storage.save_json(filename, perfil_visitante, folder_path=['data', 'perfis'])
    else:
        caminho = os.path.join("data", "perfis", filename)
        if not os.path.exists(caminho):
            os.makedirs(os.path.dirname(caminho), exist_ok=True)
            with open(caminho, "w", encoding="utf-8") as f:
                json.dump(perfil_visitante, f, indent=2, ensure_ascii=False)

def salvar_professor_config_db(professor, email, municipio, config):
    """
    Salva a configuração do professor em um arquivo JSON específico no Drive.
    Nome do arquivo: perfil_{professor_sanitized}.json
    """
    safe_name = professor.replace(" ", "_").lower()
    filename = f"perfil_{safe_name}.json"
    
    # Adiciona metadados extras
    config["email"] = email
    config["municipio"] = municipio
    
    if USE_CLOUD_STORAGE:
        google_storage.save_json(filename, config, folder_path=['data', 'perfis'])
    else:
        caminho = os.path.join("data", "perfis", filename)
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2, ensure_ascii=False)

def carregar_perfil_professor_db(nome_professor):
    """Carrega o perfil de um professor específico do arquivo JSON."""
    safe_name = nome_professor.replace(" ", "_").lower()
    filename = f"perfil_{safe_name}.json"
    
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value={}, folder_path=['data', 'perfis'])
    
    caminho = os.path.join("data", "perfis", filename)
    if os.path.exists(caminho):
        try:
            with open(caminho, "r", encoding="utf-8-sig") as f:
                content = f.read()
                if content:
                    return json.loads(content)
        except json.JSONDecodeError:
            st.warning(f"⚠️ Arquivo de perfil `{filename}` está mal formatado.")
            return {}
    return {}

def _parse_markdown_to_story(texto_markdown, styles):
    """Converte um texto em markdown simples para uma lista de 'flowables' do ReportLab."""
    story = []
    
    # Pré-processa o texto para lidar com blocos de código e listas
    lines = texto_markdown.split('\n')
    in_code_block = False
    
    for line in lines:
        stripped_line = line.strip()

        # Títulos
        if stripped_line.startswith('# '):
            story.append(Paragraph(stripped_line.replace('# ', ''), styles['H1']))
        elif stripped_line.startswith('## '):
            story.append(Paragraph(stripped_line.replace('## ', ''), styles['H2']))
        elif stripped_line.startswith('### '):
            story.append(Paragraph(stripped_line.replace('### ', ''), styles['H3']))
        
        # Linha horizontal
        elif stripped_line == '---':
            story.append(Spacer(1, 4*mm))
            # Para uma linha visual, precisaríamos de um Flowable customizado ou um Drawing.
            # Por simplicidade, usamos um parágrafo com sublinhados.
            story.append(Paragraph('_' * 80, styles['Body']))
            story.append(Spacer(1, 4*mm))

        # Citação
        elif stripped_line.startswith('> '):
            story.append(Paragraph(stripped_line.replace('> ', ''), styles['Quote']))

        # Itens de lista
        elif stripped_line.startswith(('- ', '* ')):
            # Usando um parágrafo com um marcador de emoji para melhor compatibilidade
            story.append(Paragraph(f"• {stripped_line[2:]}", styles['ListItem']))

        # Linhas com negrito e chave-valor (Ex: **Professor:** Helio)
        elif stripped_line.startswith('**') and ':**' in stripped_line:
            # Converte para <b> e deixa o Paragraph renderizar
            formatted_line = stripped_line.replace('**', '<b>', 1).replace(':**', '</b>:', 1)
            story.append(Paragraph(formatted_line, styles['Body']))

        # Parágrafo normal (com suporte a negrito no meio do texto)
        else:
            # Substitui **texto** por <b>texto</b>
            formatted_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(formatted_line, styles['Body']))
            
    return story

def gerar_pdf_aula_ia(texto_markdown):
    """Gera um PDF a partir do texto Markdown gerado pela IA, usando ReportLab."""
    buffer = io.BytesIO()
    font_family = _registrar_fontes_reportlab()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)

    # Estilos
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=10, leading=14, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=18, alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle(name='H2', fontName=f"{font_family}-Bold", fontSize=14, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name='H3', fontName=f"{font_family}-Bold", fontSize=12, spaceBefore=6, spaceAfter=3))
    styles.add(ParagraphStyle(name='Quote', parent=styles['Body'], fontName=f"{font_family}-Italic", leftIndent=15, rightIndent=15, textColor=colors.darkslategray))
    styles.add(ParagraphStyle(name='ListItem', parent=styles['Body'], leftIndent=10, firstLineIndent=-5))

    # Converte o markdown para uma lista de flowables
    story = _parse_markdown_to_story(texto_markdown, styles)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_pdf_simulado(disciplina_nome, turma_nome, professor_nome, questoes):
    """Gera um PDF de um simulado com questões e gabarito usando ReportLab."""
    buffer = io.BytesIO()
    font_family = _registrar_fontes_reportlab()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Body', fontName=font_family, fontSize=10, leading=12))
    styles.add(ParagraphStyle(name='H1', fontName=f"{font_family}-Bold", fontSize=16, alignment=TA_CENTER, spaceAfter=4))
    styles.add(ParagraphStyle(name='H2', fontName=f"{font_family}-Bold", fontSize=12, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name='Question', fontName=f"{font_family}-Bold", fontSize=11, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name='Option', parent=styles['Body'], leftIndent=10))
    styles.add(ParagraphStyle(name='GabaritoTitle', fontName=f"{font_family}-Bold", fontSize=14, alignment=TA_CENTER, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name='GabaritoItem', fontName=font_family, fontSize=12, alignment=TA_CENTER))

    story = []
    story.append(Paragraph("Simulado de Revisão", styles['H1']))
    story.append(Paragraph(f"<b>Disciplina:</b> {disciplina_nome}", styles['H2']))
    story.append(Paragraph(f"<b>Turma:</b> {turma_nome} | <b>Professor:</b> {professor_nome}", styles['H2']))
    story.append(Spacer(1, 8*mm))

    gabarito = {}
    for i, q in enumerate(questoes, 1):
        story.append(Paragraph(f"{i}. {q['question_text']}", styles['Question']))
        opcoes = q.get('options', [])
        for j, opt in enumerate(opcoes):
            letra = chr(ord('a') + j)
            story.append(Paragraph(f"{letra}) {opt}", styles['Option']))
        gabarito[i] = chr(ord('A') + q['correct_option_index'])

    story.append(PageBreak())
    story.append(Paragraph("Gabarito", styles['GabaritoTitle']))
    for num, resp in sorted(gabarito.items()):
        story.append(Paragraph(f"<b>Questão {num}:</b> {resp.upper()}", styles['GabaritoItem']))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def carregar_horario_global():
    """Carrega o quadro de horários completo da escola."""
    filename = "horario_global.json"
    if USE_CLOUD_STORAGE:
        return google_storage.load_json(filename, default_value={})
    
    caminho = os.path.join("data", filename)
    if os.path.exists(caminho):
        try:
            with open(caminho, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return {}
    return {}

def obter_horario_professor_do_global(nome_professor):
    """
    Filtra o horário global para encontrar as aulas de um professor específico.
    Retorna uma lista de dicionários formatada para DataFrame.
    """
    global_db = carregar_horario_global()
    aulas_prof = []
    
    # Mapeamento de ordem para ordenação
    ordem_dias = {"SEGUNDA- FEIRA": 1, "SEGUNDA-FEIRA": 1, "TERÇA-FEIRA": 2, "QUARTA-FEIRA": 3, "QUINTA-FEIRA": 4, "SEXTA-FEIRA": 5}
    
    for dia, periodos in global_db.items():
        for periodo, salas in periodos.items():
            for sala, dados in salas.items():
                # Verifica se o nome do professor está contido no registro (busca flexível)
                prof_db = dados.get("professor", "").lower()
                if nome_professor.lower() in prof_db or prof_db in nome_professor.lower():
                    aulas_prof.append({
                        "Dia": dia,
                        "OrdemDia": ordem_dias.get(dia, 9),
                        "Período": periodo,
                        "Horário": dados.get("horario", ""),
                        "Sala": sala,
                        "Disciplina": dados.get("disciplina", "")
                    })
    
    # Ordena por Dia e depois por Período
    aulas_prof.sort(key=lambda x: (x["OrdemDia"], x["Período"]))
    return aulas_prof

def criar_botao_voltar():
    """Cria um botão padronizado para voltar ao menu principal."""
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🏠 Voltar para o Menu Principal", use_container_width=True):
            st.switch_page("app.py")

def exibir_menu_lateral():
    """Exibe um menu de navegação lateral padronizado e botões rápidos."""
    with st.sidebar:
        # Botão explícito para voltar ao início
        st.page_link("app.py", label="Início / Menu Principal", icon="🏠", use_container_width=True)
        st.divider()

# --- FUNÇÕES DE FREQUÊNCIA ACUMULADA (POR PROFESSOR) ---

def obter_caminho_frequencia_professor(professor):
    """Retorna o caminho do arquivo de frequência acumulada do professor."""
    safe_name = professor.replace(" ", "_").lower()
    return os.path.join("data", "frequencia", f"frequencia_{safe_name}.json")

def carregar_frequencia_professor(professor):
    """Carrega o DataFrame de frequência acumulada do professor."""
    caminho = obter_caminho_frequencia_professor(professor)
    return carregar_dados_json(caminho)

def salvar_frequencia_dia(professor, turma, data_obj, df_chamada):
    """Salva/Atualiza a frequência de um dia específico no arquivo acumulado do professor."""
    # 1. Carrega dados existentes
    df_total = carregar_frequencia_professor(professor)
    
    # 2. Prepara os novos dados
    df_novos = df_chamada.copy()
    df_novos["Turma"] = turma
    df_novos["Data"] = data_obj.strftime('%Y-%m-%d')
    
    # 3. Mescla
    if df_total is None or df_total.empty:
        df_total = df_novos
    else:
        # Garante que as colunas de filtro existam
        if "Data" not in df_total.columns: df_total["Data"] = ""
        if "Turma" not in df_total.columns: df_total["Turma"] = ""
        
        # Remove registros anteriores para esta Turma e Data (Sobrescrita para atualização)
        data_str = data_obj.strftime('%Y-%m-%d')
        mask_existente = (df_total["Turma"] == turma) & (df_total["Data"] == data_str)
        df_total = df_total[~mask_existente]
        
        # Concatena
        df_total = pd.concat([df_total, df_novos], ignore_index=True)
    
    # 4. Salva
    caminho = obter_caminho_frequencia_professor(professor)
    salvar_dados_json(caminho, df_total)